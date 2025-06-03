import os
import subprocess
import zipfile
import io
from PIL import Image, ImageChops
import pytesseract
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from flask import Flask, render_template, request, send_from_directory, redirect, send_file, Response

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "static/output"

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@app.route("/", methods=["GET", "POST"])
def index():
    images = sorted([f for f in os.listdir(OUTPUT_FOLDER) if f.endswith(".png")])
    ocr_texts = {}
    ocr_tables = {}

    if request.method == "POST":
        for file in os.listdir(OUTPUT_FOLDER):
            file_path = os.path.join(OUTPUT_FOLDER, file)
            if os.path.isfile(file_path) and (file.endswith(".png") or file.endswith(".txt")):
                os.remove(file_path)

        file = request.files["pdf"]
        if file.filename.endswith(".pdf"):
            pdf_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
            file.save(pdf_path)
            convert_pdf_to_images_and_text(pdf_path)
            images = sorted([f for f in os.listdir(OUTPUT_FOLDER) if f.endswith(".png")])

    for image in images:
        txt_file = os.path.splitext(image)[0] + ".txt"
        txt_path = os.path.join(OUTPUT_FOLDER, txt_file)
        if os.path.exists(txt_path):
            with open(txt_path, "r", encoding="utf-8") as f:
                raw_text = f.read().strip()
                table = extract_table_from_text(raw_text)
                structured_pairs = [(k, v) for k, v in table if k]

                clean_lines = []
                for line in raw_text.splitlines():
                    replaced = False
                    for k, v in structured_pairs:
                        if k in line:
                            clean_lines.append(f"{k}: {v}")
                            replaced = True
                            break
                    if not replaced:
                        clean_lines.append(line)

                final_text = "\n".join(clean_lines)
                ocr_texts[image] = final_text
                ocr_tables[image] = structured_pairs

    return render_template("index.html", images=images, ocr_texts=ocr_texts, ocr_tables=ocr_tables)

def trim_image(img):
    bg = Image.new(img.mode, img.size, img.getpixel((0, 0)))
    diff = ImageChops.difference(img, bg)
    bbox = diff.getbbox()
    return img.crop(bbox) if bbox else img

def extract_table_from_text(text):
    lines = text.strip().splitlines()
    keywords = ["المؤلف", "الصفحات", "سنة النشر", "القسم"]
    table = []
    for line in lines:
        for key in keywords:
            if key in line:
                parts = line.split(key)
                value = parts[-1].strip()
                table.append((key, value))
                break
    return table

def convert_pdf_to_images_and_text(pdf_path):
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_prefix = os.path.join(OUTPUT_FOLDER, base_name)
    subprocess.run(["pdftoppm", "-png", pdf_path, output_prefix], check=True)
    for file in os.listdir(OUTPUT_FOLDER):
        if file.startswith(base_name) and file.endswith(".png"):
            image_path = os.path.join(OUTPUT_FOLDER, file)
            img = Image.open(image_path).convert("RGB")
            trimmed = trim_image(img)
            trimmed.save(image_path)
            text = pytesseract.image_to_string(trimmed, lang='ara+eng')
            txt_path = os.path.splitext(image_path)[0] + ".txt"
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(text)

@app.route("/clear", methods=["POST"])
def clear():
    for file in os.listdir(OUTPUT_FOLDER):
        file_path = os.path.join(OUTPUT_FOLDER, file)
        if os.path.isfile(file_path) and (file.endswith(".png") or file.endswith(".txt")):
            os.remove(file_path)
    return redirect("/")

@app.route("/download-all", methods=["GET"])
def download_all():
    zip_stream = io.BytesIO()
    with zipfile.ZipFile(zip_stream, "w", zipfile.ZIP_DEFLATED) as zipf:
        for filename in os.listdir(OUTPUT_FOLDER):
            if filename.endswith(".png") or filename.endswith(".txt"):
                filepath = os.path.join(OUTPUT_FOLDER, filename)
                zipf.write(filepath, arcname=filename)
    zip_stream.seek(0)
    return send_file(zip_stream, mimetype="application/zip", as_attachment=True, download_name="converted_output.zip")

@app.route("/static/output/<filename>")
def download_file(filename):
    return send_from_directory(OUTPUT_FOLDER, filename)

@app.route("/export/<filename>.docx", endpoint="export_docx")
def export_docx(filename):
    txt_name = filename + ".txt"
    txt_path = os.path.join(OUTPUT_FOLDER, txt_name)

    if not os.path.exists(txt_path):
        return "File not found", 404

    with open(txt_path, "r", encoding="utf-8") as f:
        content = f.read().strip()

    doc = Document()
    section = doc.sections[0]
    section.right_margin = Pt(36)
    section.left_margin = Pt(36)

    table_keywords = ["المؤلف", "الصفحات", "سنة النشر", "القسم"]
    table_data = extract_table_from_text(content)

    if table_data:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for k, v in table_data:
            row_cells = table.add_row().cells
            row_cells[0].text = k
            row_cells[1].text = v
            row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph("\n")

    for line in content.splitlines():
        if not any(k in line for k, _ in table_data):
            p = doc.add_paragraph()
            run = p.add_run(line)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            run.font.size = Pt(12)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return send_file(
        doc_io,
        as_attachment=True,
        download_name=f"{filename}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/export/<filename>.txt", endpoint="export_txt")
def export_txt(filename):
    txt_path = os.path.join(OUTPUT_FOLDER, f"{filename}.txt")
    if not os.path.exists(txt_path):
        return "File not found", 404

    with open(txt_path, "r", encoding="utf-8") as f:
        content = f.read()

    return Response(
        content,
        mimetype="text/plain",
        headers={"Content-Disposition": f"attachment; filename={filename}.txt"}
    )

@app.route("/export-docx", methods=["POST"])
def export_docx_all():
    doc = Document()
    section = doc.sections[0]
    section.right_margin = Pt(36)
    section.left_margin = Pt(36)

    table_keywords = ["المؤلف", "الصفحات", "سنة النشر", "القسم"]

    for filename in sorted(os.listdir(OUTPUT_FOLDER)):
        if filename.endswith(".txt"):
            base_name = os.path.splitext(filename)[0]
            txt_path = os.path.join(OUTPUT_FOLDER, filename)

            with open(txt_path, "r", encoding="utf-8") as f:
                content = f.read().strip()

            # Add page heading
            doc.add_paragraph(f"📄 Page: {base_name}", style="Heading 2")

            table_data = extract_table_from_text(content)
            if table_data:
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                for k, v in table_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = k
                    row_cells[1].text = v
                    row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            doc.add_paragraph("\n")

            for line in content.splitlines():
                if not any(k in line for k, _ in table_data):
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                    run.font.size = Pt(12)

            doc.add_paragraph("\n\n")  # Separate pages

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return send_file(
        doc_io,
        as_attachment=True,
        download_name="combined_output.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
