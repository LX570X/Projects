import os
import fitz  # PyMuPDF
import unicodedata
import re
import zipfile
from flask import Flask, request, render_template, send_file
from docx import Document
from docx.shared import Inches, Pt
from werkzeug.utils import secure_filename

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20MB max

def clean_and_format_text(text):
    cleaned = ''.join(c for c in text if unicodedata.category(c)[0] != 'C' or c == '\n')
    cleaned = cleaned.replace('\r', '').strip()
    return [p.strip() for p in re.split(r'\n{2,}', cleaned) if p.strip()]

def join_broken_lines(text):
    lines = text.split('\n')
    joined = []
    for line in lines:
        line = line.strip()
        if not line:
            joined.append("")
        elif joined and not joined[-1].endswith(('.', ':', '?', '!')) and \
             not re.match(r'^[•*0-9\-–]', line):
            joined[-1] += ' ' + line
        else:
            joined.append(line)
    return '\n'.join(joined)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        uploaded = request.files['pdf']
        if uploaded.filename == '':
            return "No file selected", 400

        filename = secure_filename(uploaded.filename)
        base = os.path.splitext(filename)[0]
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        uploaded.save(pdf_path)

        doc = fitz.open(pdf_path)
        document = Document()
        all_text = ""

        for i in range(len(doc)):
            page = doc.load_page(i)
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (round(b[1]), b[0]))
            lines = [b[4].strip() for b in blocks if b[4].strip()]
            text = join_broken_lines('\n'.join(lines))

            # add page header
            header_txt = f"========== Page {i+1} ==========\n"
            all_text += f"\n\n{header_txt}"
            hdr = document.add_paragraph(f"Page {i+1}")
            hdr.runs[0].bold = True
            hdr.paragraph_format.space_after = Pt(10)

            # add paragraphs
            for para in clean_and_format_text(text):
                all_text += "    " + para + "\n\n"
                p = document.add_paragraph(para)
                p.paragraph_format.first_line_indent = Inches(0.25)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(8)

        # write outputs
        txt_path  = os.path.join(UPLOAD_FOLDER, base + '.txt')
        docx_path = os.path.join(UPLOAD_FOLDER, base + '.docx')
        zip_path  = os.path.join(UPLOAD_FOLDER, base + '.zip')

        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(all_text.strip())
        document.save(docx_path)

        with zipfile.ZipFile(zip_path, 'w') as zf:
            zf.write(txt_path,  os.path.basename(txt_path))
            zf.write(docx_path, os.path.basename(docx_path))

        return render_template('index.html',
                               txt_file=os.path.basename(txt_path),
                               docx_file=os.path.basename(docx_path),
                               zip_file=os.path.basename(zip_path),
                               preview_text=all_text[:5000])

    return render_template('index.html')

@app.route('/download/<filename>')
def download(filename):
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename),
                     as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
