import os
import subprocess
from flask import Flask, request, render_template, send_file, redirect, url_for, flash
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from pytesseract import image_to_string
from docx.enum.text import WD_ALIGN_PARAGRAPH


UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "static/converted"
ALLOWED_EXTENSIONS = {"pdf", "png", "jpg", "jpeg", "docx", "txt", "csv", "json"}

app = Flask(__name__)
app.secret_key = "secretkey"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def convert_file(file, target_format):
    filename = secure_filename(file.filename)
    input_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(input_path)

    name, ext = os.path.splitext(filename)
    ext = ext.lower().lstrip(".")
    output_filename = f"{name}.{target_format}"
    output_path = os.path.join(OUTPUT_FOLDER, output_filename)

    try:
        # PDF → PNG
        if ext == "pdf" and target_format == "png":
            images = convert_from_path(input_path)
            if images:
                output_path = os.path.join(OUTPUT_FOLDER, f"{name}_page1.png")
                images[0].save(output_path, "PNG")

        # PNG/JPG → PDF
        elif ext in {"png", "jpg", "jpeg"} and target_format == "pdf":
            image = Image.open(input_path).convert("RGB")
            image.save(output_path, "PDF")

        # TXT → DOCX
        elif ext == "txt" and target_format == "docx":
            doc = Document()

            def is_arabic(text):
                arabic_chars = [c for c in text if '\u0600' <= c <= '\u06FF']
                return len(arabic_chars) > len(text) * 0.3  # 30% threshold

            with open(input_path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    para = doc.add_paragraph(line)
                    if is_arabic(line):
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            doc.save(output_path)

        # DOCX → TXT
        elif ext == "docx" and target_format == "txt":
            doc = Document(input_path)
            with open(output_path, "w", encoding="utf-8") as f:
                for para in doc.paragraphs:
                    f.write(para.text + "\n")
        
        # CSV → JSON
        elif ext == "csv" and target_format == "json":
            import csv, json
            with open(input_path, newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                data = list(reader)
            with open(output_path, "w", encoding="utf-8") as jsonfile:
                json.dump(data, jsonfile, indent=4, ensure_ascii=False)

# JSON → CSV
        elif ext == "json" and target_format == "csv":
            import csv, json
            with open(input_path, "r", encoding="utf-8") as jsonfile:
                data = json.load(jsonfile)
            with open(output_path, "w", newline='', encoding='utf-8') as csvfile:
                if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                    writer = csv.DictWriter(csvfile, fieldnames=data[0].keys())
                    writer.writeheader()
                    writer.writerows(data)
                else:
                    raise Exception("Unsupported JSON structure for CSV export")


        # DOCX → PDF (Linux/Mac using LibreOffice)
        elif ext == "docx" and target_format == "pdf":
            subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", OUTPUT_FOLDER,
                input_path
            ], check=True)
            output_path = os.path.join(OUTPUT_FOLDER, f"{name}.pdf")

        # PDF → DOCX (OCR-based)
        elif ext == "pdf" and target_format == "docx":
            images = convert_from_path(input_path)
            doc = Document()

            def is_arabic(text):
                arabic_chars = [c for c in text if '\u0600' <= c <= '\u06FF']
                return len(arabic_chars) > len(text) * 0.3  # if 30% or more Arabic

            for img in images:
                text = image_to_string(img, lang="ara+eng").strip()
                if not text:
                    continue
                para = doc.add_paragraph(text)
                if is_arabic(text):
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            doc.save(output_path)

        else:
            raise Exception(f"Unsupported conversion: {ext} → {target_format}")

    except Exception as e:
        print("Conversion error:", e)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"[Error during conversion: {e}]")

    return os.path.basename(output_path)

@app.route("/", methods=["GET", "POST"])
def index():
    download_url = None

    if request.method == "POST":
        if "file" not in request.files:
            flash("No file part")
            return redirect(request.url)

        file = request.files["file"]
        target_format = request.form.get("target_format")

        if file.filename == "":
            flash("No selected file")
            return redirect(request.url)

        if file and allowed_file(file.filename):
            converted_file = convert_file(file, target_format)
            download_url = url_for("download_file", filename=converted_file)
        else:
            flash("Invalid file format")

    return render_template("index.html", download_url=download_url)

@app.route("/download/<filename>")
def download_file(filename):
    path = os.path.join(app.config["OUTPUT_FOLDER"], filename)
    return send_file(path, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
